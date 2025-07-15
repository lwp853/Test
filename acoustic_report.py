"""Generate an acoustic report in Word format from a CSV file."""

import argparse
import pandas as pd
from docx import Document

# Default thresholds
RT60_LIMIT = 1.2  # seconds
STI_LIMIT = 0.6
LAEQ_LIMIT = 55  # dB(A)

def add_room_section(doc: Document, room: str, rt60: float, sti: float, laeq: float) -> None:
    """Add a formatted section for a single room."""
    doc.add_heading(room, level=1)
    doc.add_paragraph(f"RT60: {rt60:.2f} s", style="List Bullet")
    doc.add_paragraph(f"STI: {sti:.2f}", style="List Bullet")
    doc.add_paragraph(f"LAeq: {laeq:.1f} dB", style="List Bullet")

    comments = []
    if rt60 > RT60_LIMIT:
        comments.append("Consider additional absorption.")
    if sti < STI_LIMIT:
        comments.append("Speech intelligibility may be inadequate.")
    if laeq > LAEQ_LIMIT:
        comments.append("Background noise exceeds typical limits.")

    for text in comments:
        para = doc.add_paragraph(text)
        para.style = "Intense Quote"


def generate_report(csv_path: str, output_path: str) -> None:
    """Generate a Word document summarizing acoustic parameters."""
    df = pd.read_csv(csv_path)

    doc = Document()
    doc.add_heading("Acoustic Report", 0)

    for _, row in df.iterrows():
        add_room_section(
            doc,
            str(row.get("Room Name", row.get("Room", "Room"))),
            float(row.get("RT60", 0)),
            float(row.get("STI", 0)),
            float(row.get("LAeq", 0)),
        )
        doc.add_page_break()

    doc.save(output_path)


def main() -> None:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("csv", help="Path to input CSV")
    parser.add_argument("output", nargs="?", default="Acoustic_Report.docx",
                        help="Output Word file")
    args = parser.parse_args()
    generate_report(args.csv, args.output)


if __name__ == "__main__":
    main()
